"""
DOCX document loader.

Extracts text content from Microsoft Word documents using python-docx.
"""

import logging
from typing import List

from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.ingestion.loaders.base import BaseLoader, LoadedDocument

logger = logging.getLogger(__name__)


class DocxLoader(BaseLoader):
    """
    Loader for Microsoft Word documents (.docx).

    Uses python-docx for text extraction. Handles:
    - Paragraphs and headings
    - Tables (converted to text)
    - Basic metadata extraction
    """

    SUPPORTED_EXTENSIONS = [".docx"]

    def __init__(self, include_tables: bool = True):
        """
        Initialize the DOCX loader.

        Args:
            include_tables: Whether to extract text from tables
        """
        self.include_tables = include_tables

    def load(self, file_path: str) -> LoadedDocument:
        """
        Load a DOCX document.

        Args:
            file_path: Path to the DOCX file

        Returns:
            LoadedDocument with extracted text and metadata

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a DOCX
        """
        path = self._validate_file(file_path)
        logger.info(f"Loading DOCX: {path.name}")

        # Compute hash and size
        file_hash = self._compute_file_hash(file_path)
        file_size = self._get_file_size(file_path)

        # Open document
        doc = DocxDocument(file_path)

        # Extract text
        text_parts = self._extract_text(doc)
        full_text = "\n\n".join(text_parts)

        # Extract metadata
        metadata = self._extract_metadata(doc)

        return LoadedDocument(
            content=full_text,
            filename=path.name,
            source_type="docx",
            file_path=str(path.absolute()),
            file_hash=file_hash,
            file_size=file_size,
            metadata=metadata,
        )

    def _extract_text(self, doc: DocxDocument) -> List[str]:
        """
        Extract text from all document elements.

        Args:
            doc: DocxDocument instance

        Returns:
            List of text strings
        """
        text_parts = []

        # Iterate through document body elements in order
        for element in doc.element.body:
            if element.tag.endswith('p'):
                # It's a paragraph
                para = Paragraph(element, doc)
                text = para.text.strip()
                if text:
                    # Check if it's a heading
                    if para.style and para.style.name.startswith('Heading'):
                        text = f"\n{text}\n"
                    text_parts.append(text)

            elif element.tag.endswith('tbl') and self.include_tables:
                # It's a table
                table = Table(element, doc)
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(table_text)

        return text_parts

    def _extract_table_text(self, table: Table) -> str:
        """
        Convert a table to text representation.

        Args:
            table: Table object

        Returns:
            Text representation of the table
        """
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Skip empty rows
            if any(cells):
                rows.append(" | ".join(cells))

        if rows:
            return "\n".join(rows)
        return ""

    def _extract_metadata(self, doc: DocxDocument) -> dict:
        """
        Extract metadata from DOCX.

        Args:
            doc: DocxDocument instance

        Returns:
            Dictionary of metadata
        """
        metadata = {}

        try:
            core_props = doc.core_properties
            if core_props:
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.subject:
                    metadata["subject"] = core_props.subject
                if core_props.keywords:
                    metadata["keywords"] = core_props.keywords
                if core_props.created:
                    metadata["created"] = str(core_props.created)
                if core_props.modified:
                    metadata["modified"] = str(core_props.modified)
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {e}")

        # Count sections/paragraphs
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)

        return metadata

    def get_headings(self, file_path: str) -> List[str]:
        """
        Get all headings from a DOCX document.

        Args:
            file_path: Path to the DOCX file

        Returns:
            List of heading texts
        """
        doc = DocxDocument(file_path)
        headings = []

        for para in doc.paragraphs:
            if para.style and para.style.name.startswith('Heading'):
                text = para.text.strip()
                if text:
                    headings.append(text)

        return headings
